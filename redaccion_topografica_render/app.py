from flask import Flask, render_template, request, send_file, jsonify
from docx import Document
from docx.shared import Pt
import io, math, csv, re

app = Flask(__name__)

# =========================
#  NÚMEROS → TEXTO (ES)
# =========================
UNIDADES = ["cero","uno","dos","tres","cuatro","cinco","seis","siete","ocho","nueve"]
DECENAS = ["","diez","veinte","treinta","cuarenta","cincuenta","sesenta","setenta","ochenta","noventa"]
ESPECIALES_10_29 = {
    10:"diez",11:"once",12:"doce",13:"trece",14:"catorce",15:"quince",
    16:"dieciséis",17:"diecisiete",18:"dieciocho",19:"diecinueve",
    20:"veinte",21:"veintiuno",22:"veintidós",23:"veintitrés",24:"veinticuatro",25:"veinticinco",
    26:"veintiséis",27:"veintisiete",28:"veintiocho",29:"veintinueve"
}
CENTENAS = ["","cien","doscientos","trescientos","cuatrocientos","quinientos","seiscientos","setecientos","ochocientos","novecientos"]

def num_0_99(n:int)->str:
    if n<10: return UNIDADES[n]
    if 10<=n<=29: return ESPECIALES_10_29[n]
    d,u = divmod(n,10)
    return DECENAS[d] if u==0 else f"{DECENAS[d]} y {UNIDADES[u]}"

def num_0_999(n:int)->str:
    if n<100: return num_0_99(n)
    if n==100: return "cien"
    c, r = divmod(n,100)
    cent = CENTENAS[c] if c!=1 else "ciento"
    return cent if r==0 else f"{cent} {num_0_99(r)}"

def num_to_words_es(n:int)->str:
    if n<1000: return num_0_999(n)
    millones, resto = divmod(n, 1_000_000)
    miles, abajo = divmod(resto, 1000)
    parts=[]
    if millones: parts.append("un millón" if millones==1 else f"{num_0_999(millones)} millones")
    if miles: parts.append("mil" if miles==1 else f"{num_0_999(miles)} mil")
    if abajo: parts.append(num_0_999(abajo))
    return " ".join(parts) if parts else "cero"

def number_with_decimal_to_words_es(value: str) -> str:
    """
    '20.50' o '20,50' -> 'veinte punto cincuenta'
    (redondea a 2 decimales si aplica)
    """
    s = value.replace(",", ".").strip()
    if "." in s:
        f = round(float(s), 2)
        integer = int(abs(math.floor(f)))
        decimals = int(round(abs(f - integer)*100))
        sign = "-" if f<0 else ""
        int_words = num_to_words_es(integer)
        if decimals==0:
            return f"{sign}{int_words}"
        dec_words = num_to_words_es(decimals)
        return f"{sign}{int_words} punto {dec_words}"
    else:
        integer = int(s.strip())
        return num_to_words_es(integer)

# =========================
#  ESTACIONES FLEXIBLES
#  ('1' | 'A' | '1A' | '1 A')
# =========================
def parse_station_label(raw: str):
    """
    Acepta '1', 'A', '1A', '1 A'.
    Devuelve (numero:int|None, letra:str|None en mayúscula).
    """
    s = (raw or "").strip()
    m = re.match(r'^\s*(\d+)?\s*([A-Za-zÁÉÍÓÚÜÑáéíóúüñ]+)?\s*$', s)
    if not m:
        return None, None
    num = m.group(1)
    let = m.group(2)
    return (int(num) if num is not None else None,
            let.upper() if let else None)

def station_label_to_text(raw: str) -> str:
    """
    '1'   -> 'uno'
    'A'   -> 'A'
    '1A'  -> 'uno A'
    '1 A' -> 'uno A'
    (si no matchea, devuelve el raw saneado)
    """
    n, letter = parse_station_label(raw)
    parts = []
    if n is not None:
        parts.append(num_to_words_es(n))
    if letter:
        parts.append(letter)
    return " ".join(parts) if parts else (raw or "").strip()

# =========================
#  RUMBOS Y REDACCIÓN
# =========================
def rumbo_text(cardinal:str)->str:
    m = {"N":"norte","S":"sur","E":"este","W":"oeste","O":"oeste"}
    return m.get((cardinal or "").strip().upper(), (cardinal or "").lower())

def redactar_segmento(est_i_raw:str, est_f_raw:str,
                      ns:str, grados:int, minutos:int, segundos:int, ew:str,
                      distancia:str) -> str:
    # Estaciones (permitir 1 / A / 1A)
    est_i_txt = station_label_to_text(est_i_raw)
    est_f_txt = station_label_to_text(est_f_raw)

    # Distancia a palabras
    dist_txt = number_with_decimal_to_words_es(distancia)

    # Rumbos
    ns_txt = rumbo_text(ns)
    ew_txt = rumbo_text(ew)

    # Plurales
    g_txt = "grado" if grados==1 else "grados"
    m_txt = "minuto" if minutos==1 else "minutos"
    s_txt = "segundo" if segundos==1 else "segundos"

    grados_txt = num_to_words_es(grados)
    minutos_txt = num_to_words_es(minutos)
    segundos_txt = num_to_words_es(segundos)

    return (f"De la estación {est_i_txt} a la estación {est_f_txt} con una distancia de "
            f"{dist_txt} metros y un rumbo {ns_txt} {grados_txt} {g_txt}, "
            f"{minutos_txt} {m_txt} {segundos_txt} {s_txt} {ew_txt}.")

# =========================
#  RUTAS
# =========================
@app.route("/", methods=["GET"])
def index():
    return render_template("index.html")

@app.route("/preview", methods=["POST"])
def preview():
    """
    Espera líneas tipo:
    est_i, est_f, NS, grados, minutos, segundos, EW, distancia

    Ejemplos:
    1, 2, S, 46, 35, 19, E, 20.50
    A, B, N, 10, 0, 0, W, 15
    1A, 1B, S, 5, 30, 0, E, 7.25
    """
    lines = request.form.get("lineas", "").strip().splitlines()
    errores, frases = [], []
    for idx, raw in enumerate(lines, start=1):
        if not raw.strip():
            continue
        parts = [p.strip() for p in raw.split(",")]
        if len(parts) != 8:
            errores.append(f"Línea {idx}: formato inválido (se esperaban 8 campos).")
            continue
        try:
            # Estaciones ahora son strings flexibles
            est_i = parts[0]
            est_f = parts[1]
            ns = parts[2]
            grados = int(parts[3]); minutos = int(parts[4]); segundos = int(parts[5])
            ew = parts[6]
            distancia = parts[7]

            # Validaciones básicas
            if ns.upper() not in ("N","S"): raise ValueError("Rumbo NS debe ser N o S")
            if ew.upper() not in ("E","W","O"): raise ValueError("Rumbo EW debe ser E, W u O")
            if not (0 <= grados <= 359): raise ValueError("Grados fuera de rango (0-359)")
            if not (0 <= minutos < 60): raise ValueError("Minutos fuera de rango (0-59)")
            if not (0 <= segundos < 60): raise ValueError("Segundos fuera de rango (0-59)")

            frase = redactar_segmento(est_i, est_f, ns, grados, minutos, segundos, ew, distancia)
            frases.append(frase)
        except Exception as e:
            errores.append(f"Línea {idx}: {e}")
    return jsonify({"frases": frases, "errores": errores})

@app.route("/docx", methods=["POST"])
def docx():
    """
    Recibe JSON: { "frases": [...], "titulo": "..." }
    Devuelve .docx descargable
    """
    frases = request.json.get("frases", [])
    titulo = request.json.get("titulo", "Redacción de Levantamiento Topográfico")
    if not frases:
        return jsonify({"error":"No hay frases para generar."}), 400

    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)
    doc.add_heading(titulo, level=1)
    for f in frases:
        doc.add_paragraph(f)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(
        buf,
        as_attachment=True,
        download_name="redaccion_topografica.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.route("/upload_csv", methods=["POST"])
def upload_csv():
    """
    CSV con encabezados:
    est_i,est_f,NS,grados,minutos,segundos,EW,distancia
    """
    if "file" not in request.files or request.files["file"].filename == "":
        return jsonify({"error":"No se recibió archivo."}), 400

    f = request.files["file"]
    frases, errores = [], []
    try:
        data = f.read()
        try:
            text = data.decode("utf-8")
        except UnicodeDecodeError:
            text = data.decode("latin-1")
        reader = csv.DictReader(text.splitlines())
        required = ["est_i","est_f","NS","grados","minutos","segundos","EW","distancia"]
        for r in reader:
            if any(col not in r for col in required):
                return jsonify({"error":"Encabezados inválidos. Se esperaban: " + ", ".join(required)}), 400
            try:
                est_i = (r["est_i"] or "").strip()   # string flexible
                est_f = (r["est_f"] or "").strip()
                ns = (r["NS"] or "").strip()
                grados = int(r["grados"]); minutos = int(r["minutos"]); segundos = int(r["segundos"])
                ew = (r["EW"] or "").strip()
                distancia = (r["distancia"] or "").strip()

                if ns.upper() not in ("N","S"): raise ValueError("NS inválido (N/S)")
                if ew.upper() not in ("E","W","O"): raise ValueError("EW inválido (E/W/O)")
                if not (0 <= grados <= 359): raise ValueError("grados fuera de rango")
                if not (0 <= minutos < 60): raise ValueError("minutos fuera de rango")
                if not (0 <= segundos < 60): raise ValueError("segundos fuera de rango")

                frases.append(redactar_segmento(est_i, est_f, ns, grados, minutos, segundos, ew, distancia))
            except Exception as e:
                errores.append(str(e))
    except Exception as e:
        return jsonify({"error": f"No se pudo leer el CSV: {e}"}), 400

    if not frases:
        return jsonify({"error":"No se generaron frases. Revisa el CSV."}), 400

    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)
    doc.add_heading('Redacción de Levantamiento Topográfico', level=1)
    for fline in frases:
        doc.add_paragraph(fline)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(
        buf,
        as_attachment=True,
        download_name="redaccion_topografica_desde_csv.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

if __name__ == "__main__":
    # Para local; en Render se usa gunicorn (ver Procfile)
    app.run(host="0.0.0.0", port=5000, debug=True)
